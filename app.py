import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd
import datetime
import pytz
from dateutil import parser
import plotly.express as px
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import locale
from docx2pdf import convert
import tempfile
import zipfile
import os
import pythoncom  # Añadido para manejar COM

# --------------------------------------------------------------------------------
# CONFIGURACIÓN DEL LOCALE PARA MESES EN ESPAÑOL
# --------------------------------------------------------------------------------
try:
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')  # Unix/Linux
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, 'Spanish_Spain.1252')  # Windows
    except locale.Error:
        st.warning("No se pudo configurar el locale a español. Los nombres de los meses estarán en inglés.")

# --------------------------------------------------------------------------------
# CREDENCIALES DE SERVICIO
# --------------------------------------------------------------------------------
SERVICE_ACCOUNT_INFO = {
    "type": "service_account",
    "project_id": "buho-api-2024-sheets",
    "private_key_id": "49baaf5d4716afe832bb5b215ee5d783b03a5f94",
    "private_key": """-----BEGIN PRIVATE KEY-----
MIIEvQIBADANBgkqhkiG9w0BAQEFAASCBKcwggSjAgEAAoIBAQDFlko5MebPggbD
cxaiGkUUKU3bqZnlVE+vQNtOG69EuBJzUvpwQpF1TtYUnYNNyLxyZbv2kJoiywNJ
n+UaWf0tw8cPVLbKwiwPF6fKntbR42xlxiQAWspBIVVlUBKNNIsZ2R9NLavc+NLH
wQGVw8j5t3mcugASikV0PiGeMCjuLSEzOPyic4xjmTcRqeSaQ4oBOODeUkIeX6AV
ahP5lPQDpC/wvD/ZTR6U/0dStc3NJQ6he+71kXXpum9pN1C4NYAjMMAL6R/upqd3
xPxZsPwHS+pMs7cIGdjDUIIGBWMQJhoCnKqznIbxHEX7HtBv1R1oLrIIMoeFSDbX
AvUQQMlNAgMBAAECggEAEVhudyNplP1f7SAJFF10g37VwisbIr36SdMKfYqiIgj7
u6qE6D57yP3M/t3OBVHSM0O5kr1ifpvuU7QI931/Y5lce/zOmDGgDwofVYMIrj/G
CBAzIGHYAAw2VDDJlCJQ9Mmx/QM9o2YnkNghdL2NgtiHwUm10GrZiokax+mH6lKd
wTChqYPYtKttq9mREtoCGTqzPuqVR8vCdnKqAohddDfgaARNixay/PyrmqSiiWGp
bdtZleRjSL4/F7xri+WKq30NliALHBd8jzSlkSvzHLgeNR+wyRUf/GpvMZBjFG8y
B8ytXmy74h/yDzAYVj2tmgzywqiHlfLDBLVosTo1mQKBgQD2SGWcXVPp+Z0UOhjs
o52Fd4JsMK4im83D6mdxprdOO3FUBFC44HRrMBqjWqHMBHRNgHFyPN10APhSBRe8
MVxzWhXpAIM3K+5vB6hQr1PXN22Kp7bs9HGhF4vw8yWFMzUa44HlPKdVT2M6l5gn
vWe9KktbatM8pkytplnl/V8oWQKBgQDNYgmAaYssQKnyZSeru7s701SFF6O8cZqJ
VuliVwgkdcmTGTUpLZYei3QX7KHD3cAxKNZzdJdYA/G+nTkY98a5uzkd2liPmDW/
7sBRH30/H8vpKVu0YjhFqdBDsblFRRtWLebNjvs7UOcDLFfnMEkByZ3BC7utxANm
84Z6cvcKFQKBgHf4um7aU8dVjjxNNNkJtvFOT11OtXUseqbmZ+/IK+FTOZiY5Y25
4VxZuZA71TdiMBmU6S6iEaqx0kV6L57AWO3kQ2oWktTsdKDnlQmA7xGW8aiqnIR/
a17y7nu4pl1lnYf0rdEyo7z+CDOBp2Asdv2CPeVRe4c+53lr4L0VmSY5AoGAESBg
vHWQpnMJ+O2YfkicV2PLA4IyJC+w/EzkD1BEnI257mtGtJVZlFh6qNgRsTyXn0HR
iDUrvaouiX+g2EUpLCnBnIytn+PIb6XgIaOnlRD4twu82vDp0l1TwaFbWrxliC0x
tuh6aLrZWLlk5yFupRiD8CojT10uD3K1PxbBJPUCgYEA4iyM3F1LzPKl4lXsBKZD
/DC8P/mJyEONmDVGMb/T+/V4dbW327vMJP2tikLhIqAseN2/0kkd5obtnHAfUQvy
vY8iNvaf6oA0sXuCetFH6cGJoBLK6RYOOFTnwVAa7ZI6nLKVN/FHOREs4xOWYiyP
dNLv9ywuT9km02+2A/aqDDc=
-----END PRIVATE KEY-----\n""",
    "client_email": "google-api@buho-api-2024-sheets.iam.gserviceaccount.com",
    "client_id": "105531071438636619321",
    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
    "token_uri": "https://oauth2.googleapis.com/token",
    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
    "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/google-api%40buho-api-2024-sheets.iam.gserviceaccount.com",
    "universe_domain": "googleapis.com"
}

# --------------------------------------------------------------------------------
# CONFIGURACIÓN DE STREAMLIT
# --------------------------------------------------------------------------------
st.set_page_config(
    page_title="Dashboard Actividades",
    layout="wide"
)

# --------------------------------------------------------------------------------
# FUNCIONES AUXILIARES
# --------------------------------------------------------------------------------
def obtener_worksheet_df(url_hoja: str, sheet_name: str = None) -> pd.DataFrame:
    """
    Devuelve un DataFrame con los datos de una hoja específica dada la URL de Google Sheets.
    """
    # Credenciales
    creds = Credentials.from_service_account_info(
        SERVICE_ACCOUNT_INFO,
        scopes=["https://www.googleapis.com/auth/spreadsheets",
                "https://www.googleapis.com/auth/drive"]
    )
    client = gspread.authorize(creds)
    
    try:
        # Abrir Google Sheet
        doc = client.open_by_url(url_hoja)
    except Exception as e:
        st.error(f"Error al abrir la hoja de cálculo: {e}")
        return pd.DataFrame()
    
    # Seleccionar la hoja específica o la primera si no se proporciona
    try:
        if sheet_name:
            worksheet = doc.worksheet(sheet_name)
        else:
            worksheet = doc.get_worksheet(0)
    except Exception as e:
        st.error(f"Error al seleccionar la hoja '{sheet_name}': {e}")
        return pd.DataFrame()
    
    # Extraer todos los valores
    try:
        data = worksheet.get_all_values()
    except Exception as e:
        st.error(f"Error al extraer los datos de la hoja: {e}")
        return pd.DataFrame()
    
    if not data:
        st.warning("La hoja está vacía.")
        return pd.DataFrame()
    
    # Convertir a DataFrame
    headers = data[0]
    rows = data[1:]
    df = pd.DataFrame(rows, columns=headers)
    
    # Estandarizar nombres de columnas
    df.columns = (
        df.columns
        .str.strip()
        .str.lower()
        .str.replace(' ', '_', regex=False)
        .str.replace('(', '', regex=False)
        .str.replace(')', '', regex=False)
        .str.replace('-', '_', regex=False)
        .str.replace('á', 'a', regex=False)
        .str.replace('í', 'i', regex=False)
        .str.replace('é', 'e', regex=False)
        .str.replace('ó', 'o', regex=False)
        .str.replace('ú', 'u', regex=False)
    )
    
    return df

def convertir_fecha_columna(df, col):
    """
    Convierte la columna 'col' en fecha (datetime) si es posible. 
    Si no es posible, la deja como está.
    """
    try:
        df[col] = pd.to_datetime(df[col], errors='coerce', dayfirst=True)
    except:
        pass
    return df

def obtener_mes_year(df, col):
    """ Crea una nueva columna con formato 'Mes Año' basado en la columna de fecha. """
    df[f'mes_year_{col}'] = df[col].dt.strftime('%B %Y').str.capitalize()
    return df

def filtrar_por_mes_year(df, mes_year, columna_fecha="fecha_inicio"):
    """ Filtra el DF por el mes_year en la columna de fecha. Si mes_year es 'Todos', no filtra. """
    if mes_year == "Todos":
        return df
    else:
        return df[df[f'mes_year_{columna_fecha}'] == mes_year]

def calcular_actividades_concluidas(df):
    """
    Devuelve un DataFrame con las actividades concluidas y añade:
    - tiempo_real = (fecha_fin - fecha_inicio).days
    - delta = dias_estimados - tiempo_real
    """
    if 'fecha_inicio' not in df.columns:
        st.error("La columna 'fecha_inicio' no existe en el DataFrame.")
        return pd.DataFrame()

    if 'fecha_fin' not in df.columns:
        st.warning("No existe 'fecha_fin'. No se pueden calcular concluidas con exactitud.")
        return pd.DataFrame()

    # Solo las que tengan fecha_inicio y fecha_fin
    mask = df["fecha_inicio"].notnull() & df["fecha_fin"].notnull()
    concluidas = df[mask].copy()

    if concluidas.empty:
        return concluidas

    # Calcular "tiempo_real" (en días)
    concluidas["tiempo_real"] = (concluidas["fecha_fin"] - concluidas["fecha_inicio"]).dt.days

    # Convertir "dias_estimados" a numérico
    concluidas["dias_estimados"] = pd.to_numeric(concluidas["dias_estimados"], errors="coerce")

    # Calcular "delta"
    concluidas["delta"] = concluidas["dias_estimados"] - concluidas["tiempo_real"]

    # Crear columna de estado
    concluidas["estado"] = concluidas["delta"].apply(lambda x: "A tiempo" if x >= 0 else "Con retraso")

    return concluidas

def calcular_actividades_en_curso(df, tz="America/Monterrey"):
    """
    Devuelve un DataFrame con las actividades en curso y añade:
    - tiempo_real = (fecha_actual - fecha_inicio).days
    - delta = dias_estimados - tiempo_real
    """
    if 'fecha_inicio' not in df.columns or 'dias_estimados' not in df.columns:
        return pd.DataFrame()

    # Solo las que tengan fecha_inicio y NO tengan fecha_fin
    if 'fecha_fin' in df.columns:
        mask = df["fecha_inicio"].notnull() & df["fecha_fin"].isnull()
    else:
        # Si 'fecha_fin' no existe
        mask = df["fecha_inicio"].notnull()
    
    encurso = df[mask].copy()
    if encurso.empty:
        return encurso

    # Fecha de hoy
    now = datetime.datetime.now(pytz.timezone(tz)).date()

    encurso["tiempo_real"] = (now - encurso["fecha_inicio"].dt.date).apply(lambda d: d.days)
    encurso["dias_estimados"] = pd.to_numeric(encurso["dias_estimados"], errors="coerce")
    encurso["delta"] = encurso["dias_estimados"] - encurso["tiempo_real"]
    encurso["estado"] = encurso["delta"].apply(lambda x: "A tiempo" if x >= 0 else "Con retraso")

    return encurso

def calcular_actividades_pendientes(df):
    """
    Devuelve un DataFrame con las actividades pendientes (sin fecha_inicio).
    """
    if 'fecha_inicio' not in df.columns:
        st.error("La columna 'fecha_inicio' no existe en el DataFrame.")
        return pd.DataFrame()
    
    mask = df["fecha_inicio"].isnull()
    pendientes = df[mask].copy()
    pendientes["estado"] = "Pendiente"
    return pendientes

def crear_grafico_barras_estaqueadas(df):
    """
    Crea un gráfico de barras apiladas mostrando:
    - A tiempo, Con retraso, Pendiente
    por cada responsable.
    """
    if 'estado' not in df.columns or 'responsable' not in df.columns:
        return None

    if df.empty:
        return None

    data_concluidas = df[df['estado'] != 'Pendiente'].groupby(['responsable', 'estado']).size().reset_index(name='count')
    data_pendientes = df[df['estado'] == 'Pendiente'].groupby(['responsable']).size().reset_index(name='count')
    data_pendientes['estado'] = 'Pendiente'

    data = pd.concat([data_concluidas, data_pendientes], ignore_index=True)

    data_pivot = data.pivot(index='responsable', columns='estado', values='count').fillna(0).reset_index()

    estados = ['A tiempo', 'Con retraso', 'Pendiente']
    for estado in estados:
        if estado not in data_pivot.columns:
            data_pivot[estado] = 0

    data_pivot = data_pivot[['responsable'] + estados]
    data_melt = data_pivot.melt(id_vars='responsable', value_vars=estados, var_name='Estado', value_name='Cantidad')

    fig = px.bar(
        data_melt,
        x='responsable',
        y='Cantidad',
        color='Estado',
        barmode='stack',
        title='Actividades por Responsable',
        color_discrete_map={
            "A tiempo": "#66CDAA",     # Verde más oscuro
            "Con retraso": "#FF8C00",  # Naranja más fuerte
            "Pendiente": "lightgray"   # Gris claro
        },
        labels={'responsable': 'Responsable', 'Cantidad': 'Número de Actividades'}
    )

    fig.update_layout(
        xaxis_title="Responsable",
        yaxis_title="Número de Actividades",
        legend_title="Estado",
        title={
            'text': 'Actividades por Responsable',
            'y': 0.9,
            'x': 0.5,
            'xanchor': 'center',
            'yanchor': 'top'
        },
        font=dict(
            family="Arial",
            size=12,
            color="Black"
        )
    )

    return fig

def mostrar_grafico_barras_estaqueadas(df):
    fig = crear_grafico_barras_estaqueadas(df)
    if fig:
        st.plotly_chart(fig)

def eliminar_arroba_solicitante(df):
    """ Elimina el símbolo '@' al inicio de cada entrada en la columna 'solicitante'. """
    if 'solicitante' in df.columns:
        df['solicitante'] = df['solicitante'].str.lstrip('@')
    return df

def separar_responsables(df):
    """
    Separa múltiples responsables (si están separados por coma)
    y explota cada uno en una fila distinta.
    """
    if 'responsable' not in df.columns:
        return df
    
    df['responsable'] = df['responsable'].str.split(',')
    df = df.explode('responsable')
    df['responsable'] = df['responsable'].str.strip()
    return df

def preparar_datos_para_reporte(df):
    """
    Prepara los datos para generar el reporte Word en formato dict:
    {
        'concluidas': df_concluidas,
        'encurso': df_encurso,
        'pendientes': df_pendientes
    }
    """
    concluidas = calcular_actividades_concluidas(df)
    encurso = calcular_actividades_en_curso(df)
    pendientes = calcular_actividades_pendientes(df)
    return {
        'concluidas': concluidas,
        'encurso': encurso,
        'pendientes': pendientes
    }

def set_cell_background(cell, color):
    """
    Establece el color de fondo de una celda en hex (sin #).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_table_borders(table):
    """
    Aplica bordes a todas las celdas de una tabla.
    """
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')  # Tamaño del borde
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')  # Color negro
        tblBorders.append(element)
    
    tblPr.append(tblBorders)

def generar_reporte_word(area_df, responsables_df_dict, mes_str):
    """
    Genera un documento Word con la información del área y de cada responsable.
    Formato:
      - Título principal
      - Sección de Área (tabla 2 columnas: 
          [col1: gráfica, col2: tabla (Tipo/Cantidad)])
      - Sección de Responsables (por cada uno, tabla 2 columnas:
          [col1: tabla (Tipo/Cantidad), col2: gráfica])
    """
    doc = Document()

    # Fuente por defecto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Título principal
    title = f"Reporte de Actividades {mes_str}"
    paragraph = doc.add_heading(title, level=0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Eliminar espacio antes y después del párrafo
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    # ------------------------------------------------------------------
    # Sección de Área
    # ------------------------------------------------------------------
    area_heading = doc.add_heading('Reporte de Área', level=1)
    area_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    area_heading.style.font.size = Pt(16)  # Tamaño de fuente 16
    area_heading.style.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
    # Eliminar espacio antes y después del párrafo
    area_heading.paragraph_format.space_before = Pt(0)
    area_heading.paragraph_format.space_after = Pt(0)

    # Métricas de Área
    concluidas_count = len(area_df['concluidas'])
    encurso_count = len(area_df['encurso'])
    pendientes_count = len(area_df['pendientes'])

    # Crear tabla principal (1 fila, 2 columnas)
    table_area = doc.add_table(rows=1, cols=2)
    table_area.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_area.autofit = True

    # **NO** Aplicar bordes a la tabla principal
    # set_table_borders(table_area)  # Esta línea ha sido eliminada/comentada

    # Celda izquierda: gráfica (Área)
    cell_left = table_area.rows[0].cells[0]
    combined_area_df = pd.concat([
        area_df['concluidas'],
        area_df['encurso'],
        area_df['pendientes']
    ], ignore_index=True)

    fig_area = crear_grafico_barras_estaqueadas(combined_area_df)
    if fig_area:
        buf_area = BytesIO()
        fig_area.write_image(buf_area, format="png")
        buf_area.seek(0)
        # 9 cm es ~3.54 pulgadas
        cell_left.paragraphs[0].add_run().add_picture(buf_area, width=Inches(3.54))
        cell_left.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Centrar verticalmente
        cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Celda derecha: subtabla con "Tipo" y "Cantidad"
    cell_right = table_area.rows[0].cells[1]
    metrics_table = cell_right.add_table(rows=4, cols=2)
    metrics_table.autofit = True

    # Aplicar bordes solo a la subtabla de métricas
    set_table_borders(metrics_table)

    # Encabezado de la subtabla (fila 0)
    hdr_cells = metrics_table.rows[0].cells
    hdr_cells[0].text = "Tipo"
    hdr_cells[1].text = "Cantidad"

    # Dar color de fondo al header (azul) y texto en blanco
    for cell in hdr_cells:
        set_cell_background(cell, "4F81BD")  # Azul Excel
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Eliminar espacio antes y después del párrafo
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

    # Fila 1: Actividades concluidas
    row1 = metrics_table.rows[1].cells
    row1[0].text = "Actividades Concluidas"
    row1[1].text = str(concluidas_count)

    # Fila 2: Actividades en curso
    row2 = metrics_table.rows[2].cells
    row2[0].text = "Actividades en Curso"
    row2[1].text = str(encurso_count)

    # Fila 3: Actividades pendientes
    row3 = metrics_table.rows[3].cells
    row3[0].text = "Actividades Pendientes"
    row3[1].text = str(pendientes_count)

    # Dar fondo gris a las filas de datos, centrar y reducir tamaño de fuente a 10
    for row in metrics_table.rows[1:]:
        for cell in row.cells:
            set_cell_background(cell, "D3D3D3")  # Gris claro
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Centrar verticalmente
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar horizontalmente
                # Eliminar espacio antes y después del párrafo
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Tamaño de fuente reducido a 10
                    run.font.name = 'Times New Roman'

    # ------------------------------------------------------------------
    # Sección de Responsables
    # ------------------------------------------------------------------
    respons_heading = doc.add_heading('Responsables', level=1)
    respons_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    respons_heading.style.font.size = Pt(16)  # Tamaño de fuente 16
    respons_heading.style.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
    # Eliminar espacio antes y después del párrafo
    respons_heading.paragraph_format.space_before = Pt(0)
    respons_heading.paragraph_format.space_after = Pt(0)

    for resp, datos in responsables_df_dict.items():
        # Subtítulo con el nombre del responsable
        resp_heading = doc.add_heading(resp, level=2)
        resp_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        resp_heading.style.font.size = Pt(12)  # Tamaño de fuente 12
        resp_heading.style.font.color.rgb = RGBColor(0, 0, 255)  # Color azul
        # Eliminar espacio antes y después del párrafo
        resp_heading.paragraph_format.space_before = Pt(0)
        resp_heading.paragraph_format.space_after = Pt(0)

        # Métricas del responsable
        concluidas_r = len(datos['concluidas'])
        encurso_r = len(datos['encurso'])
        pendientes_r = len(datos['pendientes'])

        # Crear tabla (1 fila, 2 columnas)
        table_resp = doc.add_table(rows=1, cols=2)
        table_resp.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_resp.autofit = True

        # **NO** Aplicar bordes a la tabla principal de cada responsable
        # set_table_borders(table_resp)  # Esta línea ha sido eliminada/comentada

        # Celda izquierda: subtabla con "Tipo" y "Cantidad"
        cell_left_r = table_resp.rows[0].cells[0]
        m_table = cell_left_r.add_table(rows=4, cols=2)
        m_table.autofit = True

        # Aplicar bordes solo a la subtabla de métricas
        set_table_borders(m_table)

        # Encabezado
        hdr_r = m_table.rows[0].cells
        hdr_r[0].text = "Tipo"
        hdr_r[1].text = "Cantidad"

        # Dar color de fondo al header (azul) y texto en blanco
        for cell in hdr_r:
            set_cell_background(cell, "4F81BD")  # Azul Excel
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Eliminar espacio antes y después del párrafo
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)

        # Fila 1: Concluidas
        row1_r = m_table.rows[1].cells
        row1_r[0].text = "Actividades Concluidas"
        row1_r[1].text = str(concluidas_r)

        # Fila 2: En curso
        row2_r = m_table.rows[2].cells
        row2_r[0].text = "Actividades en Curso"
        row2_r[1].text = str(encurso_r)

        # Fila 3: Pendientes
        row3_r = m_table.rows[3].cells
        row3_r[0].text = "Actividades Pendientes"
        row3_r[1].text = str(pendientes_r)

        # Dar fondo gris a las filas de datos, centrar y reducir tamaño de fuente a 10
        for row in m_table.rows[1:]:
            for cell in row.cells:
                set_cell_background(cell, "D3D3D3")  # Gris claro
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Centrar verticalmente
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar horizontalmente
                    # Eliminar espacio antes y después del párrafo
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Tamaño de fuente reducido a 10
                        run.font.name = 'Times New Roman'

        # Celda derecha: gráfica (Responsable)
        cell_right_r = table_resp.rows[0].cells[1]
        combined_resp_df = pd.concat([
            datos['concluidas'],
            datos['encurso'],
            datos['pendientes']
        ], ignore_index=True)

        fig_resp = crear_grafico_barras_estaqueadas(combined_resp_df)
        if fig_resp:
            buf_resp = BytesIO()
            fig_resp.write_image(buf_resp, format="png")
            buf_resp.seek(0)
            # 7.5 cm es ~2.95 pulgadas
            cell_right_r.paragraphs[0].add_run().add_picture(buf_resp, width=Inches(2.95))
            cell_right_r.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Centrar verticalmente
            cell_right_r.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Espacio después de cada responsable
        # doc.add_paragraph()  # Esta línea ha sido comentada/eliminada para quitar el espacio entre responsables

    # Guardar en buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_reporte_pdf(area_df, responsables_df_dict, mes_str):
    """
    Genera un reporte PDF a partir de los datos proporcionados.
    """
    with tempfile.TemporaryDirectory() as tmpdirname:
        # Generar reporte Word
        word_buffer = generar_reporte_word(area_df, responsables_df_dict, mes_str)
        word_path = os.path.join(tmpdirname, "reporte.docx")
        with open(word_path, "wb") as f:
            f.write(word_buffer.getvalue())
        
        # Inicializar COM
        pythoncom.CoInitialize()
        
        # Convertir Word a PDF
        pdf_path = os.path.join(tmpdirname, "reporte.pdf")
        try:
            convert(word_path, pdf_path)
        except Exception as e:
            st.error(f"Error al convertir Word a PDF: {e}")
            pythoncom.CoUninitialize()
            return None
        
        # Desinicializar COM
        pythoncom.CoUninitialize()
        
        # Leer PDF
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
    
    return pdf_bytes

def generar_reportes_combinados_pdf(datos_area1, responsables_df_dict1, mes_str1,
                                   datos_area2, responsables_df_dict2, mes_str2):
    """
    Genera un archivo ZIP que contiene ambos reportes PDF (DatayMC y TecNDev).
    """
    with tempfile.TemporaryDirectory() as tmpdirname:
        # Generar reporte Word para DatayMC
        word_buffer1 = generar_reporte_word(datos_area1, responsables_df_dict1, mes_str1)
        word_path1 = os.path.join(tmpdirname, "reporte_DatayMC.docx")
        with open(word_path1, "wb") as f:
            f.write(word_buffer1.getvalue())
        
        # Inicializar COM
        pythoncom.CoInitialize()
        
        # Convertir Word a PDF para DatayMC
        pdf_path1 = os.path.join(tmpdirname, "reporte_DatayMC.pdf")
        try:
            convert(word_path1, pdf_path1)
        except Exception as e:
            st.error(f"Error al convertir Word a PDF para DatayMC: {e}")
            pythoncom.CoUninitialize()
            return None
        
        # Generar reporte Word para TecNDev
        word_buffer2 = generar_reporte_word(datos_area2, responsables_df_dict2, mes_str2)
        word_path2 = os.path.join(tmpdirname, "reporte_TecNDev.docx")
        with open(word_path2, "wb") as f:
            f.write(word_buffer2.getvalue())
        
        # Convertir Word a PDF para TecNDev
        pdf_path2 = os.path.join(tmpdirname, "reporte_TecNDev.pdf")
        try:
            convert(word_path2, pdf_path2)
        except Exception as e:
            st.error(f"Error al convertir Word a PDF para TecNDev: {e}")
            pythoncom.CoUninitialize()
            return None
        
        # Desinicializar COM
        pythoncom.CoUninitialize()
        
        # Crear archivo ZIP
        zip_path = os.path.join(tmpdirname, "reportes_combinados.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            zipf.write(pdf_path1, arcname="reporte_DatayMC.pdf")
            zipf.write(pdf_path2, arcname="reporte_TecNDev.pdf")
        
        # Leer ZIP
        with open(zip_path, "rb") as f:
            zip_bytes = f.read()
    
    return zip_bytes

# --------------------------------------------------------------------------------
# CARGA DE DATOS DESDE GOOGLE SHEETS
# --------------------------------------------------------------------------------
url_dataymc = "https://docs.google.com/spreadsheets/d/1lUg4n_agnV5ASTFSm0gRAx1z6V-qQa5pNo-QugGqBBU/edit?usp=sharing"
url_tecndev = "https://docs.google.com/spreadsheets/d/1YJ653qUFHrtHsoCKXaHaWXVCmrKvYgXkQIfh4mRbkvc/edit?usp=sharing"

df_dataymc = obtener_worksheet_df(url_dataymc)
df_tecndev = obtener_worksheet_df(url_tecndev)

# --------------------------------------------------------------------------------
# RENOMBRADO DE COLUMNAS PARA TECNDEV
# --------------------------------------------------------------------------------
tecndev_columns_mapping = {
    "tarea": "tarea",
    "proyecto": "proyecto",
    "prioridad": "prioridad",
    "solicitante": "solicitante",
    "estado": "estado",
    "devs": "responsable",
    "estimado_dias": "dias_estimados",
    "fecha_de_solicitud": "fecha_solicitud",
    "fecha_de_finalizacion": "fecha_fin",
    "fecha_de_finalización": "fecha_fin",
    "tipo": "tipo",
    "descripción": "descripcion",
    "thread_slack": "slack",
    "issue_jira": "issue_jira",
    "real_hrs": "dias_real", 
    "fecha_de_inicio": "fecha_inicio",
    "fecha_entregado": "fecha_entregado"
}
df_tecndev = df_tecndev.rename(columns=tecndev_columns_mapping)

# --------------------------------------------------------------------------------
# CONVERSIÓN DE HORAS A DÍAS PARA TECNDEV (real_hrs -> dias_real)
# --------------------------------------------------------------------------------
if 'dias_real' in df_tecndev.columns:
    df_tecndev['dias_real'] = pd.to_numeric(df_tecndev['dias_real'], errors="coerce") / 8

# --------------------------------------------------------------------------------
# SEPARACIÓN DE RESPONSABLES
# --------------------------------------------------------------------------------
df_dataymc = separar_responsables(df_dataymc)
df_tecndev = separar_responsables(df_tecndev)

# --------------------------------------------------------------------------------
# ESTANDARIZACIÓN DE FECHAS, RESPONSABLE, ETC.
# --------------------------------------------------------------------------------
cols_fecha_dataymc = ["fecha_solicitud", "fecha_inicio", "fecha_fin"]
cols_fecha_tecndev = ["fecha_solicitud", "fecha_inicio", "fecha_fin", "fecha_entregado"]

for col in cols_fecha_dataymc:
    if col in df_dataymc.columns:
        df_dataymc = convertir_fecha_columna(df_dataymc, col)

for col in cols_fecha_tecndev:
    if col in df_tecndev.columns:
        df_tecndev = convertir_fecha_columna(df_tecndev, col)

df_dataymc["responsable"] = df_dataymc["responsable"].replace(["", None], "Sin asignar")
df_tecndev["responsable"] = df_tecndev["responsable"].replace(["", None], "Sin asignar")

df_dataymc = eliminar_arroba_solicitante(df_dataymc)
df_tecndev = eliminar_arroba_solicitante(df_tecndev)

# --------------------------------------------------------------------------------
# INTERFAZ PRINCIPAL (BOTÓN SUPERIOR PARA DESCARGAR AMBOS REPORTES)
# --------------------------------------------------------------------------------
st.header("Dashboard de Actividades")

# Crear contenedor para el botón global de descarga
zip_container = st.empty()

with zip_container:
    if 'zip_reportes' not in st.session_state:
        if st.button("Generar Reportes"):
            with st.spinner("Generando ambos reportes..."):
                try:
                    # Preparar datos para reporte de Área DatayMC
                    datos_area_dataymc = preparar_datos_para_reporte(df_dataymc)
                    # Preparar datos para reporte de Área TecNDev
                    datos_area_tecndev = preparar_datos_para_reporte(df_tecndev)
                    
                    # Preparar datos para cada responsable DatayMC
                    responsables_dataymc = sorted(df_dataymc["responsable"].unique())
                    responsables_df_dict_dataymc = {}
                    for resp in responsables_dataymc:
                        df_r = df_dataymc[df_dataymc["responsable"] == resp]
                        if not df_r.empty:
                            responsables_df_dict_dataymc[resp] = preparar_datos_para_reporte(df_r)
                    
                    # Preparar datos para cada responsable TecNDev
                    responsables_tecndev = sorted(df_tecndev["responsable"].unique())
                    responsables_df_dict_tecndev = {}
                    for resp in responsables_tecndev:
                        df_r = df_tecndev[df_tecndev["responsable"] == resp]
                        if not df_r.empty:
                            responsables_df_dict_tecndev[resp] = preparar_datos_para_reporte(df_r)
                    
                    # Generar y obtener zip bytes
                    zip_bytes = generar_reportes_combinados_pdf(
                        datos_area_dataymc, responsables_df_dict_dataymc, "Todos",
                        datos_area_tecndev, responsables_df_dict_tecndev, "Todos"
                    )
                    
                    if zip_bytes:
                        st.session_state['zip_reportes'] = zip_bytes
                except Exception as e:
                    st.error(f"Error al generar los reportes combinados: {e}")
    
    # Mostrar el botón de descarga si el ZIP está generado
    if st.session_state.get('zip_reportes') is not None:
        # Reemplazar el contenido del contenedor con el botón de descarga
        zip_container.download_button(
            label="Descargar Ambos Reportes",
            data=st.session_state['zip_reportes'],
            file_name="reportes_combinados.zip",
            mime="application/zip"
        )

# --------------------------------------------------------------------------------
# INTERFAZ PRINCIPAL (TABS)
# --------------------------------------------------------------------------------
tab1, tab2 = st.tabs(["DatayMC", "TecNDev"])

# --------------------------------------------------------------------------------
# TAB 1: DATAYMC
# --------------------------------------------------------------------------------
with tab1:
    # Filtro por mes y año
    if "fecha_inicio" in df_dataymc.columns:
        df_dataymc = obtener_mes_year(df_dataymc, "fecha_inicio")
        meses_year_disponibles = sorted(df_dataymc['mes_year_fecha_inicio'].dropna().unique())
        meses_year_disponibles = ["Todos"] + meses_year_disponibles
        mes_year_seleccionado = st.selectbox("Mes y Año (DatayMC)", meses_year_disponibles, key="dataymc_mes_year")
        
        df_filtrado = filtrar_por_mes_year(df_dataymc, mes_year_seleccionado, "fecha_inicio")
        mes_str = mes_year_seleccionado
    else:
        st.error("La columna 'fecha_inicio' no existe en DatayMC.")
        df_filtrado = df_dataymc
        mes_str = "Todos"

    # Análisis: Área o Responsables
    responsables = sorted(df_filtrado["responsable"].unique())
    opciones_vista = ["Área"] + responsables
    vista_seleccionada = st.selectbox("Seleccione vista (DatayMC):", opciones_vista, key="dataymc_vista")

    if vista_seleccionada == "Área":
        # Cálculos globales
        df_concluidas = calcular_actividades_concluidas(df_filtrado)
        df_encurso = calcular_actividades_en_curso(df_filtrado)
        df_pendientes = calcular_actividades_pendientes(df_filtrado)

        combined_df = pd.concat([df_concluidas, df_encurso, df_pendientes], ignore_index=True)
        mostrar_grafico_barras_estaqueadas(combined_df)

    else:
        # Responsable específico
        df_resp = df_filtrado[df_filtrado["responsable"] == vista_seleccionada]
        df_concluidas = calcular_actividades_concluidas(df_resp)
        df_encurso = calcular_actividades_en_curso(df_resp)
        df_pendientes = calcular_actividades_pendientes(df_resp)

        combined_resp_df = pd.concat([df_concluidas, df_encurso, df_pendientes], ignore_index=True)
        mostrar_grafico_barras_estaqueadas(combined_resp_df)

    st.markdown("---")
    st.subheader("Generar Reporte (DatayMC)")

    # Preparar datos para reporte de Área
    datos_area = preparar_datos_para_reporte(df_filtrado)

    # Preparar datos para cada responsable
    responsables_df_dict = {}
    for resp in responsables:
        df_r = df_filtrado[df_filtrado["responsable"] == resp]
        if not df_r.empty:
            responsables_df_dict[resp] = preparar_datos_para_reporte(df_r)

    # Crear contenedor para el botón de descarga de DatayMC
    dataymc_container = st.empty()

    with dataymc_container:
        if 'pdf_dataymc' not in st.session_state:
            if st.button("Generar Reporte", key="generate_dataymc"):
                with st.spinner("Generando reporte DatayMC..."):
                    try:
                        pdf_bytes = generar_reporte_pdf(datos_area, responsables_df_dict, mes_str)
                        if pdf_bytes:
                            st.session_state['pdf_dataymc'] = pdf_bytes
                    except Exception as e:
                        st.error(f"Error al generar el reporte: {e}")
        
        if st.session_state.get('pdf_dataymc') is not None:
            # Reemplazar el contenido del contenedor con el botón de descarga
            dataymc_container.download_button(
                label="Descargar Reporte",
                data=st.session_state['pdf_dataymc'],
                file_name=f"reporte_actividades_{mes_str.replace(' ', '_')}_DatayMC.pdf",
                mime="application/pdf"
            )

# --------------------------------------------------------------------------------
# TAB 2: TECNDEV
# --------------------------------------------------------------------------------
with tab2:
    # Filtro por mes y año
    if "fecha_inicio" in df_tecndev.columns:
        df_tecndev = obtener_mes_year(df_tecndev, "fecha_inicio")
        meses_year_disponibles2 = sorted(df_tecndev['mes_year_fecha_inicio'].dropna().unique())
        meses_year_disponibles2 = ["Todos"] + meses_year_disponibles2
        mes_year_seleccionado2 = st.selectbox("Mes y Año (TecNDev)", meses_year_disponibles2, key="tecndev_mes_year")
        
        df_filtrado2 = filtrar_por_mes_year(df_tecndev, mes_year_seleccionado2, "fecha_inicio")
        mes_str2 = mes_year_seleccionado2
    else:
        st.error("La columna 'fecha_inicio' no existe en TecNDev.")
        df_filtrado2 = df_tecndev
        mes_str2 = "Todos"

    # Análisis: Área o Responsables
    responsables2 = sorted(df_filtrado2["responsable"].unique())
    opciones_vista2 = ["Área"] + responsables2
    vista_seleccionada2 = st.selectbox("Seleccione vista (TecNDev):", opciones_vista2, key="tecndev_vista")

    if vista_seleccionada2 == "Área":
        df_concluidas2 = calcular_actividades_concluidas(df_filtrado2)
        df_encurso2 = calcular_actividades_en_curso(df_filtrado2)
        df_pendientes2 = calcular_actividades_pendientes(df_filtrado2)

        combined_df2 = pd.concat([df_concluidas2, df_encurso2, df_pendientes2], ignore_index=True)
        mostrar_grafico_barras_estaqueadas(combined_df2)

    else:
        df_resp2 = df_filtrado2[df_filtrado2["responsable"] == vista_seleccionada2]
        df_concluidas2 = calcular_actividades_concluidas(df_resp2)
        df_encurso2 = calcular_actividades_en_curso(df_resp2)
        df_pendientes2 = calcular_actividades_pendientes(df_resp2)

        combined_resp_df2 = pd.concat([df_concluidas2, df_encurso2, df_pendientes2], ignore_index=True)
        mostrar_grafico_barras_estaqueadas(combined_resp_df2)

    st.markdown("---")
    st.subheader("Generar Reporte (TecNDev)")

    # Preparar datos para reporte de Área
    datos_area2 = preparar_datos_para_reporte(df_filtrado2)

    # Preparar datos para cada responsable
    responsables_df_dict2 = {}
    for resp in responsables2:
        df_r2 = df_filtrado2[df_filtrado2["responsable"] == resp]
        if not df_r2.empty:
            responsables_df_dict2[resp] = preparar_datos_para_reporte(df_r2)

    # Crear contenedor para el botón de descarga de TecNDev
    tecndev_container = st.empty()

    with tecndev_container:
        if 'pdf_tecndev' not in st.session_state:
            if st.button("Generar Reporte", key="generate_tecndev"):
                with st.spinner("Generando reporte TecNDev..."):
                    try:
                        pdf_bytes2 = generar_reporte_pdf(datos_area2, responsables_df_dict2, mes_str2)
                        if pdf_bytes2:
                            st.session_state['pdf_tecndev'] = pdf_bytes2
                    except Exception as e:
                        st.error(f"Error al generar el reporte (TecNDev): {e}")
        
        if st.session_state.get('pdf_tecndev') is not None:
            # Reemplazar el contenido del contenedor con el botón de descarga
            tecndev_container.download_button(
                label="Descargar Reporte",
                data=st.session_state['pdf_tecndev'],
                file_name=f"reporte_actividades_{mes_str2.replace(' ', '_')}_TecNDev.pdf",
                mime="application/pdf"
            )
